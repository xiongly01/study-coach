"""DOCX file reading utilities for AI consumption."""

from __future__ import annotations

from io import BytesIO
from typing import Any


def read_docx(file_path: str) -> str:
    """Read docx file and return markdown-formatted text.

    Args:
        file_path: Path to the .docx file

    Returns:
        Markdown-formatted string suitable for AI context
    """
    from docx import Document

    doc = Document(file_path)
    return _doc_to_markdown(doc)


def read_docx_from_bytes(data: bytes) -> str:
    """Read docx from bytes and return markdown-formatted text.

    Args:
        data: Raw bytes of a .docx file

    Returns:
        Markdown-formatted string suitable for AI context
    """
    from docx import Document

    doc = Document(BytesIO(data))
    return _doc_to_markdown(doc)


def _doc_to_markdown(doc: Any) -> str:
    """Convert python-docx Document to markdown string."""
    lines = []

    for para in doc.paragraphs:
        text = _format_paragraph(para)
        if text.strip():
            lines.append(text)

    # Convert tables
    for table in doc.tables:
        table_md = _table_to_markdown(table)
        if table_md.strip():
            lines.append(table_md)

    return "\n\n".join(lines)


def _format_paragraph(para: Any) -> str:
    """Format a paragraph with inline formatting preserved."""
    text_parts = []
    for run in para.runs:
        run_text = run.text
        if not run_text:
            continue
        if run.bold and run.italic:
            run_text = f"***{run_text}***"
        elif run.bold:
            run_text = f"**{run_text}**"
        elif run.italic:
            run_text = f"*{run_text}*"
        text_parts.append(run_text)

    text = "".join(text_parts)

    # Handle list items based on style
    style_name = para.style.name if para.style else ""
    if "List" in style_name:
        if "Number" in style_name:
            # Try to extract number from paragraph
            text = f"1. {text}"
        else:
            text = f"- {text}"

    return text


def _table_to_markdown(table: Any) -> str:
    """Convert a table to markdown format."""
    rows = []
    for i, row in enumerate(table.rows):
        cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
        rows.append("| " + " | ".join(cells) + " |")
        if i == 0:
            rows.append("| " + " | ".join(["---"] * len(cells)) + " |")
    return "\n".join(rows)
